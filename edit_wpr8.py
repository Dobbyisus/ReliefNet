from pathlib import Path

from docx import Document


INPUT_PATH = Path(r"C:\Users\Shashwat Tiwari\Desktop\MeshRelief\WPR8_signed.docx")
OUTPUT_PATH = Path(r"C:\Users\Shashwat Tiwari\Desktop\MeshRelief\WPR8_signed_edited.docx")


WORK_DONE_ITEMS = [
    "Finalized the packet structure for mesh communication with packet ID, source ID, destination ID, TTL, timestamp, and payload fields.",
    "Refactored the communication approach to use packet-based messaging instead of direct raw-string chat transmission.",
    "Outlined the MeshRouter component to handle local delivery, packet forwarding, and routing decisions within the group.",
    "Designed the SeenPacketCache logic to detect duplicate packets and avoid repeated forwarding loops.",
    "Defined TTL-based forwarding rules, including packet drop conditions and decrement-on-forward behavior.",
    "Planned flooding-based forwarding as the first routing strategy before introducing route optimization or route tables.",
    "Prepared the multi-neighbor connection model using neighbor abstractions instead of a single-peer communication flow.",
    "Mapped the required NeighborConnection and NeighborTable structures for managing multiple devices in the same group.",
    "Reviewed integration requirements for adapting the current Wi-Fi Direct communication layer to packet routing.",
    "Verified that the week 1 scope remains focused on same-group multi-hop architecture before attempting cross-group relay.",
    "Prepared the integration and validation plan for one-hop packet delivery, forwarding pipeline checks, and duplicate suppression tests.",
    "Updated the implementation roadmap to align the current ReliefNet prototype with the planned multihop architecture.",
]


def set_table_value(doc: Document, label: str, value: str) -> None:
    for table in doc.tables:
        for row in table.rows:
            first_cell = row.cells[0].text.strip()
            if first_cell == label and len(row.cells) > 1:
                row.cells[1].text = value
                return
    raise ValueError(f"Could not find table row with label: {label}")


def main() -> None:
    doc = Document(str(INPUT_PATH))

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("WPR No:"):
            paragraph.text = "WPR No: 8"
        elif text.startswith("Duration:"):
            paragraph.text = "Duration: 22/06/2026-28/06/2026"

    set_table_value(
        doc,
        "Literature review",
        (
            "Reviewed packet-based mesh communication concepts focused on same-group multihop operation, "
            "including TTL-controlled forwarding, duplicate packet suppression, flooding-based routing, "
            "neighbor discovery, and relay behavior for Android device-to-device communication."
        ),
    )

    set_table_value(
        doc,
        "Project Progress Schedule (PERT Chart)",
        (
            "Updated week 1 schedule to include: "
            "packet format definition, packet-based chat refactoring, MeshRouter design, SeenPacketCache setup, "
            "TTL forwarding rules, flooding-based routing plan, multi-neighbor abstraction design, "
            "and integration validation for same-group multihop communication."
        ),
    )

    work_done_header_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Work done in this week:":
            work_done_header_index = index
            break

    if work_done_header_index is None:
        raise ValueError("Could not find 'Work done in this week:' section")

    work_paragraphs = []
    for paragraph in doc.paragraphs[work_done_header_index + 1 :]:
        if paragraph.text.strip():
            work_paragraphs.append(paragraph)
        if len(work_paragraphs) == len(WORK_DONE_ITEMS):
            break

    if len(work_paragraphs) < len(WORK_DONE_ITEMS):
        raise ValueError("Not enough paragraphs available to replace work-done items")

    for paragraph, text in zip(work_paragraphs, WORK_DONE_ITEMS):
        paragraph.text = text

    doc.save(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
